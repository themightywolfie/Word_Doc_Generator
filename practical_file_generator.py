from tkinter import *
from tkinter import filedialog
from tkinter import scrolledtext
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tkinter


#Initialsing Document object
document = Document()

#Creating a header section
section = document.sections[0]
header = section.header
footer = section.footer

#Creating custom styles
styles = document.styles
heading = styles.add_style('Practical_Number',WD_STYLE_TYPE.PARAGRAPH)
sub_heading = styles.add_style('Sub_Head',WD_STYLE_TYPE.PARAGRAPH)
content_style = styles.add_style('Content',WD_STYLE_TYPE.PARAGRAPH)

#Formatting built-in Header style
header_style = document.styles['Header']
header_font = header_style.font
header_font.name = 'Times New Roman'
header_font.size = Pt(12)
header_font.bold =True

footer_style = document.styles['Footer']
footer_font = footer_style.font
footer_font.name = 'Times New Roman'
footer_font.size = Pt(12)
footer_font.bold = True

#Defining custom styles
head = document.styles['Practical_Number']
head_font = head.font
head_font.bold=True
head_font.name= 'Times New Roman'
head_font.size = Pt(16)

sub_head_style = document.styles['Sub_Head']
sub_head_font = sub_head_style.font
sub_head_font.name='Times New Roman'
sub_head_font.size= Pt(14)
sub_head_font.bold=True

content = document.styles['Content']
content_font = content.font
content_font.name='Times New Roman'
content_font.size = Pt(12)


window = Tk()

window.title("Welcome to File Generator")

window.geometry('650x550')

file=""
out_file=""

def choose_prog():
    global file
    file=filedialog.askopenfilename()
    prog_path_label.configure(text=file)

def choose_out():
    global out_file
    out_file=filedialog.askopenfilename()
    out_path_label.configure(text=out_file)

def generate_file():
    global file
    global out_file
    print("ID Number : ",id_entry.get())
    print("Subject : ",subj_entry.get())
    print("Heading : ",head_entry.get())
    print("Aim : ",aim_entry.get())
    print("Program Code: ",prog_entry.get("1.0",tkinter.END))
    print("Output Screenshot Path : ",out_file)
    print("Conclusion : ",concl_entry.get())
    header_data=id_entry.get()+"\t\t"+subj_entry.get()     #Insert Header Here
    footer_data = "DEPSTAR(CSE)"
    title_data=head_entry.get()            #Insert Heading Here
    aim_data=aim_entry.get()  #Insert Aim Here
    program_code_data=prog_entry.get("1.0",tkinter.END)       #Insert Code Here
    picture_path=out_file
    conclusion_data=concl_entry.get()       #Insert Conclusion Here
    #Adding header
    header_content = header.paragraphs[0]
    header_content.text = header_data
    header_content.style = document.styles["Header"]

    footer_content = footer.paragraphs[0]
    footer_content.text= footer_data
    footer_content.style = document.styles["Footer"]

    #Adding title
    title = document.add_paragraph(title_data,style='Practical_Number')
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER

    #Adding sub-title
    aim = document.add_paragraph("Aim",style='Sub_Head')

    #Adding content to sub-title
    aim_content = document.add_paragraph(aim_data,style='Content')
    aim_content.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY
    program_code = document.add_paragraph("Program Code",style='Sub_Head')


    program_code_content = document.add_paragraph(program_code_data,style='Content')
    program_code_content.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY
    output = document.add_paragraph('Output',style='Sub_Head') 
    document.add_picture(picture_path,width=Inches(3.5),height=Inches(2))

    picture = document.paragraphs[-1]
    picture.alignment= WD_ALIGN_PARAGRAPH.CENTER
    conclusion = document.add_paragraph('Conclusion',style='Sub_Head')
    conclusion_content = document.add_paragraph(conclusion_data,style='Content')
    conclusion_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    dir_ = filedialog.askdirectory()    
    document.save(dir_+"\\"+save_entry.get()+".docx")
    window.destroy()

id_label = Label(window,text="Enter your ID Number : ",anchor=W,justify=LEFT,width=30)
id_label.grid(column=0,row=0)

id_entry = Entry(window,width=50)
id_entry.grid(column=1,row=0)

Label(window,text="").grid(column=0,row=1)

subj_label = Label(window,text="Enter Subject Code and Name : ",anchor=W,justify=LEFT,width=30)
subj_label.grid(column=0,row=2)

subj_entry = Entry(window,width=50)
subj_entry.grid(column=1,row=2)

Label(window,text="").grid(column=0,row=3)

head_label = Label(window,text="Enter Heading : ",anchor=W,justify=LEFT,width=30)
head_label.grid(column=0,row=4)

head_entry = Entry(window,width=50)
head_entry.grid(column=1,row=4)

Label(window,text="").grid(column=0,row=5)

aim_label = Label(window,text="Enter Aim : ",anchor=W,justify=LEFT,width=30)
aim_label.grid(column=0,row=6)

aim_entry = Entry(window,width=50)
aim_entry.grid(column=1,row=6)

Label(window,text="").grid(column=0,row=7)

prog_label = Label(window,text="Program Code : ",anchor=W,justify=LEFT,width=30)
prog_label.grid(column=0,row=8)

prog_entry = scrolledtext.ScrolledText(window,width=30,height=10)
prog_entry.grid(column=1,row=8)


Label(window,text="").grid(column=0,row=9)

out_label = Label(window,text="Upload Output Screenshot : ",anchor=W,justify=LEFT,width=30)
out_label.grid(column=0,row=10)

out_path_label = Label(window,text="",anchor=W,justify=LEFT,width=30)
out_path_label.grid(column=1,row=10)

out_btn = Button(window,text="Choose Picture",command=choose_out)
out_btn.grid(column=2,row=10)

Label(window,text="").grid(column=0,row=11)

concl_label = Label(window,text="Enter Conclusion : ",anchor=W,justify=LEFT,width=30)
concl_label.grid(column=0,row=12)

concl_entry = Entry(window,width=50)
concl_entry.grid(column=1,row=12)

Label(window,text="").grid(column=0,row=13)

save_label = Label(window,text="Save with name : ",anchor=W,justify=LEFT,width=30)
save_label.grid(column=0,row=14)

save_entry = Entry(window,width=50)
save_entry.grid(column=1,row=14)

Label(window,text="").grid(column=0,row=15)

gen_btn= Button(window,text="Generate File",command=generate_file)
gen_btn.grid(column=1,row=16)
window.mainloop()
